import cv2
import os
from sklearn.cluster import KMeans
import json
from docx import Document
import numpy as np

# Função para aplicar o algoritmo k-médias e reconstruir a imagem
def aplicar_kmeans_e_reconstruir(imagem, k):
    # Redimensiona a imagem para um vetor de pixels
    pixels = imagem.reshape((-1, 3))
    # Aplica k-médias
    kmeans = KMeans(n_clusters=k, n_init='auto')
    # Treinamento do modelo
    kmeans.fit(pixels)
    # Rótulos atribuídos a cada pixel
    rotulos = kmeans.labels_
    # Centroides
    centroides = kmeans.cluster_centers_
    # Reconstrói a imagem usando os centroides
    imagem_reconstruida = centroides[rotulos].reshape(imagem.shape)

    return imagem_reconstruida, centroides

def getInfoOfFile(path: str) -> {"resolution": tuple, "size": str, "colors-quantity": int}:
    image = cv2.imread(path)
    resolution = (image.shape[1], image.shape[0])  # (w, h)
    size_kb = os.path.getsize(path) / 1024  # Size in KB
    unique_colors = len(np.unique(image.reshape((-1, 3)), axis=0))
    return {'resolution': resolution, 'size': f"{size_kb:.2f} KB", 'colors-quantity': unique_colors}

#kValues = [2, 5, 10, 15, 20, 25, 30]
kValues = [2, 10, 30, 50, 70, 90, 128]
pasta_originais = "originais/"
pasta_destino = ""

# Busca imagens na pasta
imagens = sorted(os.listdir(pasta_originais))

# Criação do documento Word
doc = Document()

for imagem in imagens:
    dadosImagem = []
    imagemOriginal = cv2.imread(pasta_originais+imagem)
    info_original = getInfoOfFile(pasta_originais+imagem)
    dadosImagem.append({"Imagem": imagem[0:-4], "K": "Original", **info_original})

    # Criação da tabela de imagens no documento Word
    table_imagens = doc.add_table(rows=2, cols=4)  # 2 linhas, 5 colunas
    cell_original = table_imagens.cell(0, 0)
    
    # Adiciona o texto "Original" acima da imagem original
    cell_original.text = "Original"
    
    # Adiciona a imagem original à célula mantendo as proporções
    cell_original.add_paragraph().add_run().add_picture(
        os.path.join(pasta_originais, imagem),
        width=cell_original.width+10
    )

    # Loop sobre cada valor de k
    for col, k in enumerate(kValues, 1):
        # Aplica k-médias e reconstrói a imagem
        imagem_reconstruida, _ = aplicar_kmeans_e_reconstruir(imagemOriginal, k)
        # Cria o diretório se não existir | retirar .jpg do nome da pasta
        pasta_k = f'k-means/{imagem[0:-4]}/k{k}'
        if not os.path.exists(pasta_k):
            os.makedirs(pasta_k)
        # Obtém o nome da imagem sem o caminho
        nome_imagem_sem_caminho = os.path.basename(imagem)

        # Salva a imagem na pasta correspondente ao valor de k
        caminho_imagem_reconstruida = os.path.join(pasta_k, f'k{k}_reconstruida_{nome_imagem_sem_caminho}')
        cv2.imwrite(caminho_imagem_reconstruida, imagem_reconstruida)

        info_reconstruida = getInfoOfFile(caminho_imagem_reconstruida)
        name_json = f'{pasta_k}/dados_{imagem[0:-4]}.json'
        #salva Json
        with open(r''+name_json, 'w') as json_file:
            json.dump({
                            "imagem": imagem[0:-4],
                            "k":k,
                            "original":info_original,
                            "nova":info_reconstruida
            }, json_file, indent=4)

        #construir tabelas
        dadosImagem.append({"Imagem": imagem[0:-4], "K": k, **info_reconstruida})
        # Adiciona a imagem reconstruída à célula mantendo as proporções
        cell = table_imagens.cell(0, col)
        cell.text = f"K = {k}"
        cell.paragraphs[0].add_run().add_picture(
            os.path.join(pasta_destino, f'k-means/{imagem[0:-4]}/k{k}/k{k}_reconstruida_{imagem}'),
            width=cell.width
        )
    
    # Adiciona uma nova linha para a tabela de dados
    doc.add_paragraph("\n")
    
    # Criação da tabela de dados no documento Word
    table_dados = doc.add_table(rows=len(dadosImagem) + 1, cols=5)  # ajustado o número de colunas
    
    # Adiciona cabeçalhos à tabela de dados
    for col, header in enumerate(['Imagem', 'K', 'resolution', 'size', 'colors-quantity']):
        table_dados.cell(0, col).text = header
    
    # Adiciona os dados da imagem à tabela de dados
    for row, dado in enumerate(dadosImagem, 1):
        for col, valor in enumerate([dado['Imagem'], dado['K'], dado['resolution'], dado['size'], dado['colors-quantity']]):
            table_dados.cell(row, col).text = str(valor)

    # Adiciona uma nova linha para a proxima tabela
    doc.add_paragraph("\n")

# Salva o documento Word
doc.save('resultados.docx')
